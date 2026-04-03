from flask import Flask, render_template, request, jsonify
from docx import Document
import re
from datetime import datetime

app = Flask(__name__)

def detect_document_type(doc):
    """Detect if document is birthdays or service anniversaries"""
    for para in doc.paragraphs:
        text = para.text.strip()
        if para.runs and para.runs[0].bold:
            # Check for service anniversary pattern (e.g., "1 year", "5 years")
            if re.match(r'^\d+\s+years?$', text):
                return 'service'
            # Check for birthday pattern (month and day)
            if re.match(r'^[A-Z][a-z]+\s+\d+$', text):
                return 'birthday'
    return 'unknown'

def parse_birthday_document(doc):
    """Parse birthday document and extract data"""
    dates_data = {}
    current_date = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # Check if this is a date header (bold text with date format)
        if para.runs and para.runs[0].bold:
            # Extract date from bold text like "September 8"
            date_match = re.match(r'\**(.*?)\**

def parse_service_document(doc):
    """Parse service anniversary document and extract data"""
    sections = {}
    current_section = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # Check if this is a year header (bold text like "1 year" or "5 years")
        if para.runs and para.runs[0].bold:
            year_match = re.match(r'^(\d+)\s+years?$', text)
            if year_match:
                current_section = text
                sections[current_section] = []
        elif current_section:
            # This is a name under the current section
            sections[current_section].append(text)
    
    return sections

def split_cols(names):
    """Split names into 3 columns for service anniversaries"""
    n = len(names)
    cols = [[], [], []]
    per_col = (n + 2) // 3
    for i, name in enumerate(names):
        cols[min(i // per_col, 2)].append(name)
    return cols

def generate_birthday_html(dates_data):
    """Generate Bootstrap HTML from birthday/date data"""
    html_parts = []
    dates = list(dates_data.keys())
    
    # Process dates in groups of 4 (col-md-3 means 4 columns per row)
    for i in range(0, len(dates), 4):
        html_parts.append('<div class="row">')
        
        # Get up to 4 dates for this row
        row_dates = dates[i:i+4]
        
        for date in row_dates:
            names = dates_data[date]
            names_html = ' <br/> '.join(names)
            
            html_parts.append(f'''  <div class="col-md-3">
    <h3>{date}<br/></h3>
    <p>{names_html}</p>
  </div>''')
        
        html_parts.append('</div>')
    
    return '\n'.join(html_parts)

def generate_service_html(sections):
    """Generate Bootstrap HTML from service anniversary data"""
    html_parts = []
    
    for section, names in sections.items():
        cols = split_cols(names)
        html_parts.append(f'<p>\n   <strong>{section}</strong></p>')
        html_parts.append('<div class="row">')
        
        for col in cols:
            html_parts.append('   <div class="col-md-4">')
            if col:
                html_parts.append('      <p>' + '<br/> '.join(col) + '</p>')
            else:
                html_parts.append('      <p>&#160;</p>')
            html_parts.append('   </div>')
        
        html_parts.append('</div>')
    
    return '\n'.join(html_parts)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not file.filename.endswith('.docx'):
        return jsonify({'error': 'Please upload a .docx file'}), 400
    
    try:
        # Parse the document
        doc = Document(file)
        
        # Detect document type
        doc_type = detect_document_type(doc)
        
        # Reset file pointer for re-reading
        file.seek(0)
        doc = Document(file)
        
        if doc_type == 'birthday':
            data = parse_birthday_document(doc)
            html_output = generate_birthday_html(data)
            dates = list(data.keys())
            title = f"{dates[0]} - {dates[-1]} Birthdays" if dates else "Birthdays"
        elif doc_type == 'service':
            data = parse_service_document(doc)
            html_output = generate_service_html(data)
            title = "Service Anniversaries"
        else:
            return jsonify({'error': 'Unable to detect document type. Please ensure dates or year headers are in bold.'}), 400
        
        return jsonify({
            'success': True,
            'html': html_output,
            'title': title,
            'type': doc_type
        })
    
    except Exception as e:
        return jsonify({'error': f'Error processing file: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True), text)
            if date_match:
                current_date = date_match.group(1).strip()
                dates_data[current_date] = []
        elif current_date:
            # This is a name under the current date
            dates_data[current_date].append(text)
    
    return dates_data

def parse_service_document(doc):
    """Parse service anniversary document and extract data"""
    sections = {}
    current_section = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # Check if this is a year header (bold text like "1 year" or "5 years")
        if para.runs and para.runs[0].bold:
            year_match = re.match(r'^(\d+)\s+years?$', text)
            if year_match:
                current_section = text
                sections[current_section] = []
        elif current_section:
            # This is a name under the current section
            sections[current_section].append(text)
    
    return sections

def split_cols(names):
    """Split names into 3 columns for service anniversaries"""
    n = len(names)
    cols = [[], [], []]
    per_col = (n + 2) // 3
    for i, name in enumerate(names):
        cols[min(i // per_col, 2)].append(name)
    return cols

def generate_birthday_html(birthdays):
    """Generate Bootstrap HTML from birthday data"""
    html_parts = []
    dates = list(birthdays.keys())
    
    # Process dates in groups of 4 (col-md-3 means 4 columns per row)
    for i in range(0, len(dates), 4):
        html_parts.append('<div class="row">')
        
        # Get up to 4 dates for this row
        row_dates = dates[i:i+4]
        
        for date in row_dates:
            names = birthdays[date]
            names_html = ' <br/> '.join(names)
            
            html_parts.append(f'''  <div class="col-md-3">
    <h3>{date}<br/></h3>
    <p>{names_html}</p>
  </div>''')
        
        html_parts.append('</div>')
    
    return '\n'.join(html_parts)

def generate_service_html(sections):
    """Generate Bootstrap HTML from service anniversary data"""
    html_parts = []
    
    for section, names in sections.items():
        cols = split_cols(names)
        html_parts.append(f'<p>\n   <strong>{section}</strong></p>')
        html_parts.append('<div class="row">')
        
        for col in cols:
            html_parts.append('   <div class="col-md-4">')
            if col:
                html_parts.append('      <p>' + '<br/> '.join(col) + '</p>')
            else:
                html_parts.append('      <p>&#160;</p>')
            html_parts.append('   </div>')
        
        html_parts.append('</div>')
    
    return '\n'.join(html_parts)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not file.filename.endswith('.docx'):
        return jsonify({'error': 'Please upload a .docx file'}), 400
    
    try:
        # Parse the document
        doc = Document(file)
        
        # Detect document type
        doc_type = detect_document_type(doc)
        
        # Reset file pointer for re-reading
        file.seek(0)
        doc = Document(file)
        
        if doc_type == 'birthday':
            data = parse_birthday_document(doc)
            html_output = generate_birthday_html(data)
            dates = list(data.keys())
            title = f"{dates[0]} - {dates[-1]} Birthdays" if dates else "Birthdays"
        elif doc_type == 'service':
            data = parse_service_document(doc)
            html_output = generate_service_html(data)
            title = "Service Anniversaries"
        else:
            return jsonify({'error': 'Unable to detect document type. Please ensure dates or year headers are in bold.'}), 400
        
        return jsonify({
            'success': True,
            'html': html_output,
            'title': title,
            'type': doc_type
        })
    
    except Exception as e:
        return jsonify({'error': f'Error processing file: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True)
